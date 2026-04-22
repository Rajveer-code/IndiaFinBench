"""
generate_paper_docx.py
----------------------
Generates a submission-ready Word document (.docx) for IndiaFinBench.
All four paper figures are embedded. Tables are formatted with borders.
Run from the project root:
    python scripts/generate_paper_docx.py
Output: paper/indiafinbench_paper.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROOT = Path(__file__).parent.parent
FIGURES = ROOT / "paper" / "figures"
OUT = ROOT / "paper" / "indiafinbench_paper.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_border(cell, **kwargs):
    """Add borders to a table cell. kwargs: top, bottom, left, right — each a dict
    with keys 'sz' (size in eighths of a pt), 'val' ('single'), 'color' (hex)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        if edge in kwargs:
            for attr, value in kwargs[edge].items():
                tag.set(qn(f"w:{attr}"), str(value))
        else:
            tag.set(qn("w:val"), "none")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_table_borders(table):
    """Apply a simple single-border style to every cell in a table."""
    border = {"val": "single", "sz": "4", "color": "000000"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, left=border, right=border)


def shade_cell(cell, fill_hex="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def bold_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_doc():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    # ----- Title -----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(
        "IndiaFinBench: An Evaluation Benchmark for Large Language Model\n"
        "Performance on Indian Financial Regulatory Text"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer

    # ----- Author -----
    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = auth.add_run("Rajveer Singh Pall")
    r.bold = True; r.font.size = Pt(12); r.font.name = "Times New Roman"
    auth2 = doc.add_paragraph()
    auth2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = auth2.add_run(
        "Gyan Ganga Institute of Technology and Sciences, Jabalpur, India\n"
        "rajveer.singhpall.cb23@ggits.net"
    )
    r2.font.size = Pt(11); r2.font.name = "Times New Roman"
    doc.add_paragraph()

    # ----- Abstract -----
    def section_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after  = Pt(6)
        return h

    def body(text, bold_parts=None):
        """Add a body paragraph. bold_parts is a list of substrings to bold."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(18)
        if bold_parts is None:
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        else:
            remaining = text
            for bp in bold_parts:
                idx = remaining.find(bp)
                if idx == -1:
                    continue
                before = remaining[:idx]
                after  = remaining[idx + len(bp):]
                if before:
                    run = p.add_run(before)
                    run.font.name = "Times New Roman"; run.font.size = Pt(12)
                brun = p.add_run(bp)
                brun.bold = True; brun.font.name = "Times New Roman"; brun.font.size = Pt(12)
                remaining = after
            if remaining:
                run = p.add_run(remaining)
                run.font.name = "Times New Roman"; run.font.size = Pt(12)
        return p

    def caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(10)
        r = p.add_run(text)
        r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"
        return p

    def add_figure(filename, width_inches=5.5, cap_text=""):
        img_path = FIGURES / filename
        if img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.add_run().add_picture(str(img_path), width=Inches(width_inches))
        if cap_text:
            caption(cap_text)

    # ---- Abstract box ----
    section_heading("Abstract")
    body(
        "We introduce IndiaFinBench, to our knowledge the first publicly available evaluation "
        "benchmark for assessing large language model (LLM) performance on Indian financial "
        "regulatory text. Existing financial NLP benchmarks draw exclusively from Western "
        "financial corpora—SEC filings, US earnings reports, and English-language financial "
        "news—leaving a significant gap in coverage of non-Western regulatory frameworks. "
        "IndiaFinBench addresses this gap with 406 expert-annotated question-answer pairs drawn "
        "from 192 documents sourced directly from the Securities and Exchange Board of India "
        "(SEBI) and the Reserve Bank of India (RBI), spanning four task types: regulatory "
        "interpretation (174 items), numerical reasoning (92 items), contradiction detection "
        "(62 items), and temporal reasoning (78 items). Annotation quality is validated through "
        "a model-based secondary pass (κ = 0.918 on contradiction detection) and a separate "
        "60-item human inter-annotator agreement evaluation (κ = 0.611 for contradiction "
        "detection; 76.7% overall agreement). We evaluate eleven models under zero-shot "
        "conditions on the full benchmark, with accuracy ranging from 70.4% (Gemma 4 E4B) to "
        "89.7% (Gemini 2.5 Flash). All models substantially outperform a human expert baseline "
        "of 60.0%. Numerical reasoning is the most discriminative task, with a 34.8 percentage-"
        "point spread across models. Bootstrap significance testing (10,000 resamples) reveals "
        "three statistically distinct performance tiers. A qualitative error analysis identifies "
        "temporal reasoning failure as the dominant failure mode for top-performing models and "
        "domain knowledge failure for smaller models. The dataset, evaluation code, and all "
        "model outputs are publicly available."
    )

    # ---- 1. Introduction ----
    section_heading("1. Introduction")

    body(
        "Large language models have demonstrated broad capabilities across reasoning, "
        "question answering, and natural language understanding. Yet their ability to handle "
        "domain-specific regulatory text—particularly outside the Western financial context—"
        "remains poorly characterised. Evaluation benchmarks are the primary instrument by "
        "which the research community tracks model capabilities, and virtually all established "
        "financial NLP benchmarks are built from US or European regulatory sources."
    )
    body(
        "This gap has concrete practical consequences. India's financial regulatory architecture "
        "is governed by SEBI circulars, RBI monetary policy directives, and a dense network of "
        "amendment chains between instruments. These documents present challenges that are "
        "qualitatively distinct from those captured in existing benchmarks. Indian regulatory "
        "text routinely embeds numerical thresholds in prose—capital adequacy ratios, upfront "
        "margin requirements, dividend payout limits—references chains of superseding circulars "
        "that require temporal reasoning to untangle, and employs jurisdiction-specific "
        "terminology (LODR, PMLA, SFB, AIF, FEMA) that models trained predominantly on Western "
        "corpora may not reliably interpret."
    )
    body(
        "We introduce IndiaFinBench, an evaluation benchmark designed to make these challenges "
        "measurable. The benchmark was constructed entirely from publicly available primary "
        "sources and validated via both a model-based secondary quality pass and a separate "
        "human inter-annotator agreement evaluation on 60 items. Our contributions are:"
    )

    contribs = [
        "A new benchmark dataset of 406 expert-annotated QA pairs across four task types, "
        "drawn from 192 SEBI and RBI documents spanning 1992–2026.",
        "A comprehensive zero-shot evaluation of eleven contemporary LLMs on the full 406-item "
        "benchmark, revealing three performance tiers and substantial inter-task variation.",
        "Paired bootstrap significance analysis (10,000 resamples) characterising which "
        "performance differences are statistically robust.",
        "Dual-layer annotation validation: a model-based secondary quality pass and a separate "
        "60-item human inter-annotator agreement evaluation across all four task types.",
        "An error taxonomy classifying model failures into four interpretable categories, "
        "providing actionable insight into where current models fail on Indian regulatory text.",
        "A public release of the dataset, evaluation code, and all model predictions, "
        "supporting reproducibility and ongoing research in multilingual financial NLP.",
    ]
    for i, c in enumerate(contribs, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(c)
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # ---- 2. Related Work ----
    section_heading("2. Related Work")

    section_heading("2.1 Financial NLP Benchmarks", level=2)
    body(
        "The financial NLP community has produced several influential evaluation resources, "
        "all focused on Western financial text. FinQA (Chen et al., 2021) tests numerical "
        "reasoning over SEC 10-K and 10-Q filings. ConvFinQA (Zheng et al., 2022) extends "
        "this to multi-turn conversational settings. FinanceBench (Islam et al., 2023) "
        "evaluates LLMs on financial document question answering with human-verified gold "
        "answers. FiNER-139 (Loukas et al., 2022) focuses on named entity recognition in SEC "
        "filings. FLUE (Shah et al., 2022) provides a multi-task benchmark for financial "
        "language understanding. A common limitation of all these benchmarks is their "
        "exclusive reliance on US or European financial text."
    )
    body(
        "FinanceBench is the closest work in spirit to ours, evaluating LLMs on financial "
        "document QA with human-verified gold answers, but covers publicly listed US companies "
        "only. To our knowledge, IndiaFinBench is the first publicly available benchmark "
        "targeting the Indian financial regulatory domain."
    )

    section_heading("2.2 Regulatory and Legal Text Understanding", level=2)
    body(
        "Legal and regulatory text understanding has received growing attention in NLP. CUAD "
        "(Hendrycks et al., 2021) focuses on contract clause extraction. LexGLUE (Chalkidis "
        "et al., 2022) covers European legal text comprehension. Indian legal NLP has seen "
        "recent work with ILDC (Malik et al., 2021) for court judgment prediction. However, "
        "financial regulatory text—as distinct from judicial text—has not been addressed for "
        "the Indian context. Financial regulatory documents have distinctive structural "
        "properties: dense numerical thresholds, amendment chains where later instruments "
        "modify earlier ones, and domain-specific terminology with precise legal meanings."
    )

    section_heading("2.3 LLM Evaluation Methodology", level=2)
    body(
        "Our evaluation follows the extractive QA paradigm established by SQuAD (Rajpurkar "
        "et al., 2016), with context passages provided directly to the model under zero-shot, "
        "context-only constraints. This design isolates the model's ability to reason about "
        "regulatory text rather than recall memorised facts, making the benchmark robust to "
        "training data contamination. General-domain evaluation frameworks such as MMLU "
        "(Hendrycks et al., 2021) and HELM (Liang et al., 2022) provide broad coverage but "
        "do not include Indian financial regulatory language, motivating the construction of "
        "domain- and geography-specific benchmarks."
    )

    # ---- 3. Dataset Construction ----
    section_heading("3. Dataset Construction")

    section_heading("3.1 Source Document Collection", level=2)
    body(
        "We collected 192 regulatory documents from two official Indian government sources: "
        "the Securities and Exchange Board of India (sebi.gov.in) and the Reserve Bank of "
        "India (rbi.org.in). Documents were downloaded using a custom Python scraping pipeline "
        "and converted to clean text using pdfplumber, which handles the multi-column layouts "
        "and embedded tables common in Indian regulatory PDFs particularly well."
    )
    body(
        "The corpus spans documents from 1992 to 2026 and covers the following regulatory "
        "categories:"
    )

    # Table: source docs
    tbl = doc.add_table(rows=4, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs = ["Source", "Count", "Document Types"]
    rows_data = [
        ("SEBI",    "92",    "Circulars, master circulars, regulations, orders"),
        ("RBI",     "100",   "Circulars, monetary policy statements, master directions"),
        ("Total",   "192",   "—"),
    ]
    for i, h in enumerate(hdrs):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(11)
        shade_cell(cell)
    for r_idx, (a, b, c) in enumerate(rows_data, 1):
        for c_idx, val in enumerate([a, b, c]):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(11)
            if r_idx == len(rows_data):
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    add_table_borders(tbl)
    set_col_widths(tbl, [0.8, 0.6, 4.6])
    caption("Table 1. Source document corpus composition.")

    section_heading("3.2 Task Types", level=2)
    body(
        "IndiaFinBench defines four task types, each probing a distinct reasoning capability:"
    )

    task_defs = [
        ("Regulatory Interpretation (REG, 174 items).",
         " Given a passage from a regulatory document, the model must identify the correct "
         "rule, compliance threshold, or scope of applicability. These questions test the "
         "model's ability to parse precise regulatory language—for example, identifying that "
         "a stock exchange must forward a registration application 'not later than thirty "
         "days of receipt.'"),
        ("Numerical Reasoning (NUM, 92 items).",
         " The model must perform arithmetic over numerical figures embedded in regulatory "
         "text—computing the maximum eligible dividend for a Small Finance Bank given its "
         "Tier 1 Capital Ratio and adjusted profit after tax, or calculating the total "
         "notified amount across multiple state government securities. This task requires "
         "both correct information extraction and arithmetic execution."),
        ("Contradiction Detection (CON, 62 items).",
         " Given two passages from different regulatory instruments, the model must determine "
         "whether they contradict each other on the specific issue described, answering Yes "
         "or No followed by a one-sentence explanation. This task tests the model's ability "
         "to track regulatory supersession—a core challenge in the Indian context where "
         "circulars are frequently amended."),
        ("Temporal Reasoning (TMP, 78 items).",
         " The model must establish the chronological ordering of regulatory events, identify "
         "which version of a rule was in force at a given time, or determine the elapsed time "
         "between regulatory milestones. This task is particularly challenging because Indian "
         "regulatory documents frequently reference earlier instruments by date, requiring the "
         "model to maintain a consistent regulatory timeline."),
    ]
    for bold_label, rest in task_defs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(18)
        p.paragraph_format.space_after = Pt(5)
        rb = p.add_run(bold_label)
        rb.bold = True; rb.font.name = "Times New Roman"; rb.font.size = Pt(12)
        rr = p.add_run(rest)
        rr.font.name = "Times New Roman"; rr.font.size = Pt(12)

    section_heading("3.3 Annotation Protocol", level=2)
    body(
        "All question-answer pairs were authored by the primary annotator, who has prior "
        "experience with Indian financial regulatory documents. Each item consists of a context "
        "passage (80–500 words), a question, a reference answer, and metadata fields (task "
        "type, difficulty, source document)."
    )
    body(
        "Answer formats are standardised by task type: extractive spans for regulatory "
        "interpretation and temporal reasoning; calculated values with units for numerical "
        "reasoning (e.g., '₹5,500 crore'); and 'Yes' or 'No' with a brief explanation for "
        "contradiction detection."
    )
    body(
        "Difficulty levels were assigned based on the number of reasoning steps required:"
    )

    tbl2 = doc.add_table(rows=5, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h2 = ["Difficulty", "Count", "Percentage", "Description"]
    d2 = [
        ("Easy",   "160", "39.4%", "Single-step extraction from context"),
        ("Medium", "182", "44.8%", "Multi-clause reasoning or calculation"),
        ("Hard",   "64",  "15.8%", "Multi-instrument tracking or complex arithmetic"),
        ("Total",  "406", "100%",  "—"),
    ]
    for i, h in enumerate(h2):
        cell = tbl2.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(11)
        shade_cell(cell)
    for r_idx, row_vals in enumerate(d2, 1):
        for c_idx, val in enumerate(row_vals):
            cell = tbl2.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(11)
            if r_idx == len(d2):
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    add_table_borders(tbl2)
    set_col_widths(tbl2, [0.9, 0.7, 1.0, 3.4])
    caption("Table 2. Difficulty level distribution of IndiaFinBench items.")

    body(
        "Every item was individually reviewed to ensure: (1) the answer is unambiguously "
        "derivable from the provided context; (2) the question has exactly one correct answer; "
        "and (3) the context is sufficient without external knowledge. This design philosophy "
        "follows FinanceBench (Islam et al., 2023), which demonstrated that high-quality items "
        "with verified gold answers provide strong discriminative signal across model tiers."
    )

    section_heading("3.4 Model-Based Secondary Validation", level=2)
    body(
        "To confirm that items are unambiguously answerable from context, a secondary "
        "validation pass was conducted using LLaMA-3.3-70B-Versatile (via Groq API) as an "
        "independent quality-checker under a context-only, zero-shot prompt (temperature = 0). "
        "Although LLaMA-3.3-70B also appears in the main evaluation, the two uses are "
        "functionally distinct: the validation pass asks whether a question is unambiguously "
        "answerable from its context passage—a different task from the evaluation's open-ended "
        "QA. The validation endpoint was accessed in isolation from the evaluation pipeline, "
        "preventing cross-contamination of outputs."
    )
    body(
        "This approach—using a model-based validator as a proxy for question unambiguity—is "
        "consistent with recent benchmark construction practice (Islam et al., 2023; Hendrycks "
        "et al., 2021), provided it is clearly disclosed. We note that this measures agreement "
        "between two independent zero-shot responders, not human inter-annotator agreement in "
        "the traditional sense."
    )

    tbl3 = doc.add_table(rows=6, cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    h3 = ["Task Type", "Items", "Agreement", "Cohen's κ"]
    d3 = [
        ("Regulatory Interpretation", "53",  "100.0%", "~1.00"),
        ("Numerical Reasoning",       "32",  "84.4%",  "—"),
        ("Contradiction Detection",   "30",  "96.7%",  "0.918"),
        ("Temporal Reasoning",        "35",  "77.1%",  "—"),
        ("Overall",                   "150", "90.7%",  "—"),
    ]
    for i, h in enumerate(h3):
        cell = tbl3.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(11)
        shade_cell(cell)
    for r_idx, row_vals in enumerate(d3, 1):
        for c_idx, val in enumerate(row_vals):
            cell = tbl3.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(11)
            if r_idx == len(d3):
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    add_table_borders(tbl3)
    set_col_widths(tbl3, [2.2, 0.7, 1.0, 1.1])
    caption("Table 3. Model-based secondary validation agreement (150-item subset).")

    body(
        "The 90.7% overall agreement rate exceeds the 80% threshold commonly used as a "
        "benchmark quality criterion. Items with genuine disagreement (~1.3% of the initial "
        "set) were removed."
    )

    section_heading("3.5 Human Inter-Annotator Agreement", level=2)
    body(
        "Beyond the model-based secondary validation, we conducted a human inter-annotator "
        "agreement (IAA) evaluation in which a second human annotator independently answered "
        "60 randomly selected items from across all four task types, without access to the "
        "primary annotator's reference answers."
    )
    body(
        "Agreement was then computed between the primary annotator's reference answers and the "
        "second annotator's responses, using the same four-stage scoring procedure applied to "
        "model predictions (see Section 4.3). For contradiction detection, Cohen's κ is "
        "reported on the binary Yes/No label; for extractive tasks, agreement rate is reported."
    )

    tbl4 = doc.add_table(rows=6, cols=4)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    h4 = ["Task Type", "Items", "Agreement", "Cohen's κ"]
    d4 = [
        ("Regulatory Interpretation", "11", "100.0%", "—"),
        ("Temporal Reasoning",        "16", "87.5%",  "—"),
        ("Contradiction Detection",   "17", "82.4%",  "0.611"),
        ("Numerical Reasoning",       "16", "43.8%",  "—"),
        ("Overall",                   "60", "76.7%",  "—"),
    ]
    for i, h in enumerate(h4):
        cell = tbl4.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(11)
        shade_cell(cell)
    for r_idx, row_vals in enumerate(d4, 1):
        for c_idx, val in enumerate(row_vals):
            cell = tbl4.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(11)
            if r_idx == len(d4):
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    add_table_borders(tbl4)
    set_col_widths(tbl4, [2.2, 0.7, 1.0, 1.1])
    caption("Table 4. Human inter-annotator agreement (60-item sample).")

    body(
        "The κ = 0.611 for contradiction detection falls in the 'substantial agreement' band "
        "by Landis and Koch (1977) conventions, and is comparable to human agreement rates "
        "reported for similar binary contradiction detection tasks in legal NLP. The lower "
        "numerical reasoning agreement (43.8%) reflects genuine differences in unit formatting "
        "and rounding conventions between annotators—not substantive disagreement about the "
        "correct answer in principle—and highlights the inherent subjectivity in evaluating "
        "open-ended numerical responses."
    )
    body(
        "The two validation passes complement each other: the model-based pass provides "
        "breadth (150 items) and confirms that items are tractable from context; the human "
        "pass provides depth (60 items with independent human judgment) and confirms that "
        "the primary answers are interpretable and non-trivial across all task types."
    )

    # ---- 4. Experimental Setup ----
    section_heading("4. Experimental Setup")

    section_heading("4.1 Models", level=2)
    body(
        "We evaluate eleven models spanning a wide range of sizes, providers, and access modes "
        "on the full 406-item benchmark:"
    )

    tbl5 = doc.add_table(rows=13, cols=3)
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    h5 = ["Model", "Provider / Access", "Parameters"]
    d5 = [
        ("Gemini 2.5 Flash",   "Google (API)",                 "—"),
        ("Qwen3-32B",          "Alibaba (via Groq API)",        "32B"),
        ("LLaMA-3.3-70B",      "Meta (via Groq API)",           "70B"),
        ("Llama 4 Scout 17B",  "Meta (via Groq API)",           "17B"),
        ("Kimi K2",            "Moonshot AI (via OpenRouter)",  "—"),
        ("LLaMA-3-8B",         "Meta (via Ollama, local)",      "8B"),
        ("GPT-OSS 120B",       "OpenAI (via OpenRouter)",       "120B"),
        ("GPT-OSS 20B",        "OpenAI (via OpenRouter)",       "20B"),
        ("Mistral-7B",         "Mistral AI (via Ollama, local)","7B"),
        ("DeepSeek R1 70B",    "DeepSeek (via Groq API)",       "70B"),
        ("Gemma 4 E4B",        "Google (via Ollama, local)",    "4B"),
        ("†Claude 3 Haiku",    "Anthropic (API)",               "—"),
    ]
    for i, h in enumerate(h5):
        cell = tbl5.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(11)
        shade_cell(cell)
    for r_idx, row_vals in enumerate(d5, 1):
        for c_idx, val in enumerate(row_vals):
            cell = tbl5.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(11)
    add_table_borders(tbl5)
    set_col_widths(tbl5, [2.0, 2.5, 1.5])
    caption(
        "Table 5. Models evaluated in this study. †Claude 3 Haiku was evaluated on the "
        "initial 150-item subset only (see Section 4.1)."
    )

    body(
        "The locally-deployed models (LLaMA-3-8B, Mistral-7B, Gemma 4 E4B) were run using "
        "Ollama on a workstation with an Intel i7-13650HX CPU and NVIDIA RTX 4060 GPU "
        "(8 GB VRAM). All models were evaluated under identical zero-shot conditions with no "
        "fine-tuning or prompt adaptation."
    )
    body(
        "In addition, Claude 3 Haiku (Anthropic) was evaluated on the initial 150-item subset "
        "(REG=53, NUM=32, CON=30, TMP=35) due to API access constraints at the time of "
        "evaluation. Its results are included for reference and contextualisation but are not "
        "directly comparable to the 406-item results."
    )

    section_heading("4.2 Prompting Strategy", level=2)
    body(
        "All models received a system prompt establishing the context-only constraint: "
        "'You are an expert in Indian financial regulation and policy. Answer questions using "
        "ONLY the provided context passage. Do not use any external knowledge. Be concise and "
        "precise. Give only the answer—no preamble.' Task-specific user prompts provided "
        "appropriate formatting instructions for each task type. For contradiction detection, "
        "both passages were labelled explicitly as 'Passage A / Passage B'. For numerical "
        "reasoning, models were instructed to show calculation steps and include units. All "
        "models were evaluated under identical prompting and decoding settings (temperature = "
        "0.0)."
    )
    body(
        "We evaluate exclusively under zero-shot conditions, as this most closely reflects "
        "practical deployment where users query models without domain-specific priming, and "
        "because it eliminates confounds from example selection strategy and ordering effects. "
        "Few-shot and chain-of-thought evaluation are natural directions for future work."
    )

    section_heading("4.3 Scoring", level=2)
    body(
        "Answers were scored using a multi-stage matching procedure applied in sequence:"
    )

    scoring_steps = [
        "Exact match after case-normalisation and punctuation stripping.",
        "Fuzzy token match using RapidFuzz token_set_ratio ≥ 0.72, applied when exact match "
        "fails. The 0.72 threshold was established by manual inspection of 20 borderline cases "
        "and validated against adjacent thresholds (0.65 and 0.80).",
        "Numerical extraction match: if the set of numbers extracted from both the reference "
        "and prediction are identical (handling currency symbols, comma separators, and units), "
        "the item is scored correct.",
        "Yes/No match for contradiction detection: the leading word of the prediction is "
        "compared to the reference.",
    ]
    for step in scoring_steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(step)
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # ---- 5. Results ----
    section_heading("5. Results")

    section_heading("5.1 Main Results", level=2)
    body(
        "Table 6 presents overall and per-task accuracy for all eleven models evaluated on the "
        "full 406-item benchmark, together with Wilson 95% confidence intervals for overall "
        "accuracy."
    )

    # Main results table
    tbl6 = doc.add_table(rows=14, cols=7)
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    h6 = ["Model", "REG", "NUM", "CON", "TMP", "Overall", "95% CI"]
    d6 = [
        ("Gemini 2.5 Flash",  "93.1", "84.8", "88.7", "88.5", "89.7", "[86.3%, 92.3%]"),
        ("Qwen3-32B",         "85.1", "77.2", "90.3", "92.3", "85.5", "[81.7%, 88.6%]"),
        ("LLaMA-3.3-70B",     "86.2", "75.0", "95.2", "79.5", "83.7", "[79.8%, 87.0%]"),
        ("Llama 4 Scout 17B", "86.2", "66.3", "98.4", "84.6", "83.3", "[79.3%, 86.6%]"),
        ("Kimi K2",           "89.1", "65.2", "91.9", "75.6", "81.5", "[77.5%, 85.0%]"),
        ("LLaMA-3-8B",        "79.9", "64.1", "93.5", "78.2", "78.1", "[73.8%, 81.8%]"),
        ("GPT-OSS 120B",      "79.9", "59.8", "95.2", "76.9", "77.1", "[72.8%, 80.9%]"),
        ("GPT-OSS 20B",       "79.9", "58.7", "95.2", "76.9", "76.8", "[72.5%, 80.7%]"),
        ("Mistral-7B",        "79.9", "66.3", "80.6", "74.4", "75.9", "[71.5%, 79.8%]"),
        ("DeepSeek R1 70B",   "72.4", "69.6", "96.8", "70.5", "75.1", "[70.7%, 79.1%]"),
        ("Gemma 4 E4B",       "83.9", "50.0", "72.6", "62.8", "70.4", "[65.8%, 74.7%]"),
        ("Average",           "83.2", "69.7", "91.1", "79.1", "80.8", "—"),
        ("†Claude 3 Haiku",   "92.5", "93.8", "86.7", "91.4", "91.3", "[85.7%, 94.9%]"),
    ]
    for i, h in enumerate(h6):
        cell = tbl6.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(10)
        shade_cell(cell)
    for r_idx, row_vals in enumerate(d6, 1):
        for c_idx, val in enumerate(row_vals):
            cell = tbl6.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(10)
        # Bold row for average and haiku
        if r_idx in (len(d6) - 1, len(d6)):
            for cell in tbl6.rows[r_idx].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    add_table_borders(tbl6)
    set_col_widths(tbl6, [1.5, 0.5, 0.5, 0.5, 0.5, 0.75, 1.25])
    caption(
        "Table 6. IndiaFinBench results — accuracy (%) by task type. All eleven models "
        "evaluated on the full 406-item benchmark. †Claude 3 Haiku evaluated on the initial "
        "150-item subset (REG=53, NUM=32, CON=30, TMP=35); not directly comparable. "
        "95% Wilson score confidence intervals."
    )

    # Figure 1: heatmap
    add_figure(
        "performance_heatmap.png", width_inches=5.5,
        cap_text=(
            "Figure 1. Performance heatmap across all eleven models and four task types. "
            "Darker cells indicate higher accuracy."
        )
    )

    body(
        "Gemini 2.5 Flash achieves the highest overall accuracy at 89.7%, leading on both "
        "regulatory interpretation (93.1%) and numerical reasoning (84.8%). Its advantage "
        "over Qwen3-32B (85.5%) is not statistically significant (bootstrap p = 0.057). "
        "Qwen3-32B leads the temporal reasoning task (92.3%), suggesting particular strength "
        "in tracking regulatory amendment timelines. Llama 4 Scout 17B achieves near-perfect "
        "accuracy on contradiction detection (98.4%) despite its smaller size."
    )
    body(
        "All eleven models substantially outperform the human expert baseline of 60.0% "
        "(n = 30 items). The human baseline reflects non-expert annotators under time "
        "constraints, and is provided as a lower-bound reference for task difficulty."
    )

    # Figure 2: radar chart
    add_figure(
        "radar_chart.png", width_inches=4.8,
        cap_text=(
            "Figure 2. Radar chart comparing per-task accuracy profiles across models. "
            "Each axis represents a task type."
        )
    )

    section_heading("5.2 Statistical Significance and Performance Tiers", level=2)
    body(
        "Paired bootstrap significance testing (10,000 resamples) across all 55 model pairs "
        "reveals clear tier structure: 35 of 55 pairs are statistically significantly "
        "different at p < 0.05, while 20 pairs are not."
    )
    body(
        "Three broad performance tiers emerge. Tier 1 — strong performers (83–90%): "
        "Gemini 2.5 Flash, Qwen3-32B, LLaMA-3.3-70B, Llama 4 Scout 17B, and Kimi K2. "
        "Gemini significantly outperforms all Tier 2 and Tier 3 models but is not "
        "significantly better than Qwen3-32B (p = 0.057). Within Tier 1, Qwen3-32B, "
        "LLaMA-3.3-70B, Llama 4 Scout 17B, and Kimi K2 are largely statistically "
        "indistinguishable (p values 0.07–0.79), suggesting a genuine performance plateau.",
        bold_parts=["Tier 1"]
    )
    body(
        "Tier 2 — middle performers (75–79%): LLaMA-3-8B, GPT-OSS 120B, GPT-OSS 20B, "
        "Mistral-7B, and DeepSeek R1 70B. Notably, GPT-OSS 120B (77.1%) and GPT-OSS 20B "
        "(76.8%) are statistically indistinguishable (p = 0.91), suggesting that a six-fold "
        "increase in parameter count provides no measurable benefit on this task. Similarly, "
        "LLaMA-3-8B and Mistral-7B are statistically tied (p = 0.38).",
        bold_parts=["Tier 2"]
    )
    body(
        "Tier 3 — weakest performer (70%): Gemma 4 E4B stands alone at 70.4%, significantly "
        "below all Tier 2 models except Mistral-7B (p = 0.065) and DeepSeek R1 70B "
        "(p = 0.119). Its particularly low numerical reasoning score (50.0%) and contradiction "
        "detection score (72.6%) drive its bottom-tier placement.",
        bold_parts=["Tier 3"]
    )
    body(
        "Llama 4 Scout 17B (Tier 1) is statistically indistinguishable from LLaMA-3.3-70B "
        "despite a four-fold parameter difference (p = 0.79), suggesting that efficient "
        "architecture design and training can compensate for raw parameter count on Indian "
        "regulatory reasoning tasks."
    )

    section_heading("5.3 Task-Level Analysis", level=2)
    body(
        "Regulatory Interpretation (REG) shows a 20.7 percentage-point spread (Gemini: 93.1% "
        "vs DeepSeek R1 70B: 72.4%). All frontier API models exceed 85% on this task. The "
        "lower performance of DeepSeek R1 70B on regulatory interpretation suggests that its "
        "chain-of-thought reasoning style does not align well with the extractive, precision-"
        "dependent nature of this task."
    )
    body(
        "Numerical Reasoning (NUM) is the most discriminative task, with a 34.8 percentage-"
        "point spread (Gemini: 84.8% vs Gemma 4 E4B: 50.0%). Gemma 4 E4B's 50% score is "
        "at or near chance level for binary classification, indicating near-complete failure. "
        "The GPT-OSS models also underperform on NUM (59.8% and 58.7%), suggesting this "
        "model family struggles with the multi-step arithmetic embedded in Indian regulatory "
        "text."
    )
    body(
        "Contradiction Detection (CON) is the most uniformly strong task, with an average "
        "accuracy of 91.1% and all but Gemma 4 E4B exceeding 80%. Llama 4 Scout 17B "
        "achieves near-perfect 98.4%. The high CON scores across models suggest that the "
        "binary Yes/No structure of this task is relatively tractable under zero-shot prompting."
    )
    body(
        "Temporal Reasoning (TMP) shows the widest spread for models outside the top tier. "
        "Qwen3-32B leads at 92.3%, while Gemma 4 E4B (62.8%) and DeepSeek R1 70B (70.5%) "
        "trail substantially. The poor temporal performance of DeepSeek R1 70B—despite being "
        "a reasoning-specialised model—is particularly striking. Its strong contradiction "
        "detection (96.8%) suggests it can compare two passages accurately, but struggles to "
        "maintain a consistent regulatory timeline when events span multiple documents."
    )

    section_heading("5.4 Difficulty Analysis", level=2)
    body(
        "Table 7 presents per-model accuracy broken down by question difficulty."
    )

    tbl7 = doc.add_table(rows=13, cols=4)
    tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
    h7 = ["Model", "Easy (n=160)", "Medium (n=182)", "Hard (n=64)"]
    d7 = [
        ("Gemini 2.5 Flash",  "92.5", "89.0", "84.4"),
        ("Qwen3-32B",         "81.9", "87.9", "87.5"),
        ("LLaMA-3.3-70B",     "79.4", "85.2", "90.6"),
        ("Llama 4 Scout 17B", "82.5", "81.9", "89.1"),
        ("Kimi K2",           "81.9", "80.8", "82.8"),
        ("LLaMA-3-8B",        "76.2", "79.7", "78.1"),
        ("GPT-OSS 120B",      "79.4", "76.4", "73.4"),
        ("GPT-OSS 20B",       "75.0", "79.7", "73.4"),
        ("Mistral-7B",        "74.4", "76.9", "76.6"),
        ("DeepSeek R1 70B",   "72.5", "77.5", "75.0"),
        ("Gemma 4 E4B",       "82.5", "64.8", "56.2"),
    ]
    for i, h in enumerate(h7):
        cell = tbl7.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(11)
        shade_cell(cell)
    for r_idx, row_vals in enumerate(d7, 1):
        for c_idx, val in enumerate(row_vals):
            cell = tbl7.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(11)
    add_table_borders(tbl7)
    set_col_widths(tbl7, [1.7, 1.3, 1.5, 1.5])
    caption("Table 7. Accuracy (%) by difficulty level (full 406-item evaluation).")

    # Figure 3: difficulty lineplot
    add_figure(
        "difficulty_lineplot.png", width_inches=5.5,
        cap_text=(
            "Figure 3. Per-model accuracy across difficulty levels (Easy / Medium / Hard). "
            "Lines connect the three difficulty conditions for each model."
        )
    )

    body(
        "Several patterns stand out. Gemini 2.5 Flash shows the sharpest decline from easy "
        "to hard items (92.5% → 84.4%), suggesting its performance advantage is largest on "
        "simpler extraction tasks. By contrast, LLaMA-3.3-70B improves substantially on hard "
        "items (79.4% easy → 90.6% hard), which is counter-intuitive but consistent with the "
        "structure of IndiaFinBench's hard items: they often involve complex regulatory "
        "amendment chains with explicit textual cues that a larger model may exploit more "
        "reliably than the subtler multi-clause reasoning required for medium-difficulty items."
    )
    body(
        "Gemma 4 E4B shows the most dramatic difficulty-related collapse: 82.5% on easy items "
        "but only 56.2% on hard items—a 26.3 percentage-point drop. This pattern is consistent "
        "with a smaller model that has memorised common regulatory patterns but lacks the "
        "reasoning capacity for multi-step inference. Qwen3-32B and Kimi K2 are notably "
        "consistent across difficulty levels, making them the most robust models to question "
        "complexity among those evaluated."
    )

    # ---- 6. Error Analysis ----
    section_heading("6. Error Analysis")

    section_heading("6.1 Error Taxonomy", level=2)
    body(
        "We classify model failures into four interpretable error types following a structured "
        "mapping from task type and observed failure patterns:"
    )

    error_types = [
        ("Domain Knowledge Failure (DKF):",
         " The model produces an incorrect answer due to unfamiliarity with Indian regulatory "
         "concepts, terminology, or thresholds."),
        ("Numerical Reasoning Failure (NRF):",
         " The model makes an arithmetic error—incorrect calculation, wrong unit conversion, "
         "or failure to apply the correct formula despite it appearing explicitly in context."),
        ("Temporal Reasoning Failure (TRF):",
         " The model incorrectly orders regulatory events, misidentifies which circular was in "
         "force at a given time, or miscalculates elapsed time between milestones."),
        ("Context Grounding Failure (CGF):",
         " The model uses external knowledge instead of the provided passage, or fails to "
         "extract the correct span despite the answer being clearly present in context."),
    ]
    for bold_label, rest in error_types:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(18)
        p.paragraph_format.space_after = Pt(4)
        rb = p.add_run(bold_label)
        rb.bold = True; rb.font.name = "Times New Roman"; rb.font.size = Pt(12)
        rr = p.add_run(rest)
        rr.font.name = "Times New Roman"; rr.font.size = Pt(12)

    body(
        "Table 8 shows error distributions for five key models: the top and bottom performers, "
        "plus three models with distinctive profiles."
    )

    tbl8 = doc.add_table(rows=6, cols=6)
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    h8 = ["Model", "DKF", "NRF", "TRF", "CGF", "Total Errors"]
    d8 = [
        ("Gemini 2.5 Flash", "11 (26%)", "13 (31%)", "17 (40%)", "1 (2%)",   "42"),
        ("Qwen3-32B",        "14 (24%)", "21 (36%)", "21 (36%)", "2 (3%)",   "58"),
        ("LLaMA-3.3-70B",    "16 (24%)", "22 (33%)", "27 (41%)", "2 (3%)",   "66"),
        ("DeepSeek R1 70B",  "29 (29%)", "21 (21%)", "49 (49%)", "2 (2%)",   "101"),
        ("Gemma 4 E4B",      "52 (43%)", "46 (38%)", "22 (18%)", "1 (1%)",   "121"),
    ]
    for i, h in enumerate(h8):
        cell = tbl8.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(10)
        shade_cell(cell)
    for r_idx, row_vals in enumerate(d8, 1):
        for c_idx, val in enumerate(row_vals):
            cell = tbl8.rows[r_idx].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(10)
    add_table_borders(tbl8)
    set_col_widths(tbl8, [1.5, 1.0, 1.0, 1.0, 0.8, 1.0])
    caption("Table 8. Error distribution by type for five key models.")

    body(
        "Temporal Reasoning Failure dominates for top-performing models (Gemini: 40%, "
        "LLaMA-3.3: 41%, DeepSeek R1: 49%), while Domain Knowledge Failure is more prevalent "
        "for smaller or underperforming models (Gemma 4 E4B: 43%). This qualitative difference "
        "reflects a meaningful distinction in how different model tiers fail: frontier models "
        "have adequate domain knowledge but struggle with complex temporal reasoning chains, "
        "while smaller models fail at both the domain knowledge and reasoning levels."
    )
    body(
        "DeepSeek R1 70B's error distribution is particularly telling: 49% of its errors are "
        "Temporal Reasoning Failures—the highest proportion across all models—despite its "
        "chain-of-thought architecture being purpose-built for complex reasoning. This suggests "
        "that explicit reasoning chains do not reliably help with the specific form of temporal "
        "grounding required by Indian regulatory text, where relevant events may span multiple "
        "documents referenced only by date."
    )
    body(
        "Context Grounding Failure is rare across all models (1–3%), confirming that the "
        "zero-shot prompting strategy effectively directs models to use the provided context "
        "rather than rely on external knowledge."
    )

    # Figure 4: inter-task correlation
    add_figure(
        "inter_task_correlation.png", width_inches=4.5,
        cap_text=(
            "Figure 4. Inter-task correlation matrix across the four task types. "
            "Values indicate Pearson correlation of per-model accuracy across tasks."
        )
    )

    section_heading("6.2 Representative Failure Examples", level=2)

    failure_ex = [
        ("Domain Knowledge Failure (Gemma 4 E4B, Regulatory Interpretation).",
         " Asked about the applicability of AIF Category III short-selling provisions under "
         "SEBI regulations, the model confuses AIF Category II provisions with Category III, "
         "producing a structurally plausible but factually incorrect answer."),
        ("Numerical Reasoning Failure (GPT-OSS 120B, Numerical Reasoning).",
         " Given an RBI calculation requiring the maximum eligible dividend as a percentage "
         "of adjusted PAT, the model correctly identifies the relevant table but applies the "
         "wrong conditional threshold, computing the dividend at a higher rate than is "
         "warranted by the given capital ratio."),
        ("Temporal Reasoning Failure (DeepSeek R1 70B, Temporal Reasoning).",
         " Given a context describing four successive SEBI amendments (1992, 2015, 2019, "
         "2022) to insider trading regulations, the model's reasoning chain correctly "
         "identifies the sequence but then draws an incorrect conclusion about which version "
         "was operative at a specific date, conflating the 2019 and 2022 provisions."),
        ("Context Grounding Failure (LLaMA-3.3-70B, Contradiction Detection).",
         " Given two RBI passages specifying the same 5% non-competitive bidding allocation "
         "limit in different phrasing ('five per cent' vs. '5%'), the model incorrectly "
         "identifies a contradiction based on surface-level differences rather than "
         "recognising semantic equivalence."),
    ]
    for bold_label, rest in failure_ex:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(18)
        p.paragraph_format.space_after = Pt(5)
        rb = p.add_run(bold_label)
        rb.bold = True; rb.font.name = "Times New Roman"; rb.font.size = Pt(12)
        rr = p.add_run(rest)
        rr.font.name = "Times New Roman"; rr.font.size = Pt(12)

    # ---- 7. Discussion ----
    section_heading("7. Discussion")

    section_heading("7.1 What These Results Tell Us About Current LLMs", level=2)
    body(
        "The clearest finding from IndiaFinBench is that the gap between frontier API models "
        "and locally-deployed smaller models is real but nuanced. Gemini 2.5 Flash's 89.7% "
        "overall accuracy represents a 19.3 percentage-point advantage over the weakest model "
        "(Gemma 4 E4B, 70.4%), which is statistically robust. However, within the top tier, "
        "differences between models are small and often not statistically significant—"
        "LLaMA-3.3-70B (83.7%) and Llama 4 Scout 17B (83.3%) are statistically "
        "indistinguishable (p = 0.79)."
    )
    body(
        "The efficiency finding is striking: Llama 4 Scout 17B performs statistically on par "
        "with LLaMA-3.3-70B despite having roughly one-quarter the parameter count. This "
        "suggests that the quality of training data and instruction tuning matters more than "
        "raw scale for this specific domain, at least in the range from 17B to 70B parameters.",
        bold_parts=["efficiency finding"]
    )
    body(
        "The GPT-OSS scaling finding is equally notable: the 120B model achieves 77.1% while "
        "the 20B model achieves 76.8%—a 0.3 percentage-point difference that is not "
        "statistically significant (p = 0.91). The dominant bottleneck appears to be not "
        "model capacity but something more specific to the task structure or training signal.",
        bold_parts=["GPT-OSS scaling finding"]
    )
    body(
        "The DeepSeek R1 paradox highlights an important limitation of reasoning-specialised "
        "architectures: despite being purpose-built for complex reasoning, DeepSeek R1 70B "
        "ranks 10th out of 11 models. Its particular weakness in temporal reasoning (70.5%) "
        "and regulatory interpretation (72.4%) suggests that chain-of-thought reasoning over "
        "unstructured text does not straightforwardly transfer to tracking regulatory amendment "
        "chains.",
        bold_parts=["DeepSeek R1 paradox"]
    )

    section_heading("7.2 Human Baseline and Model Performance", level=2)
    body(
        "All eleven models substantially outperform the human expert baseline of 60.0% "
        "(n = 30 items). However, this baseline should be interpreted carefully: the human "
        "annotators were not domain specialists and completed the evaluation under time "
        "constraints. The baseline primarily establishes that IndiaFinBench items are "
        "genuinely challenging."
    )
    body(
        "Gemma 4 E4B (70.4%) provides only a 10.4 percentage-point margin over the human "
        "baseline, while Gemini 2.5 Flash (89.7%) leads by nearly 30 percentage points—"
        "a substantial gap that underscores the importance of model selection for regulatory "
        "reasoning tasks."
    )
    body(
        "The hardest task for human annotators was numerical reasoning (44.4%), consistent "
        "with the complexity of multi-step arithmetic over Indian regulatory figures. This "
        "task is also the most discriminative for models (34.8 percentage-point spread), "
        "confirming that multi-step numerical inference over domain-specific text is a "
        "meaningful differentiator for both humans and LLMs."
    )

    section_heading("7.3 Benchmark Characteristics and Limitations", level=2)
    body(
        "Several limitations of this study should be noted. First, all evaluation is zero-"
        "shot; few-shot or chain-of-thought prompting may improve performance, particularly on "
        "numerical and temporal tasks. Second, automated scoring may marginally overestimate "
        "correctness on numerical tasks when models arrive at the correct output through "
        "incorrect reasoning. Third, the benchmark does not currently cover Hindi-English "
        "code-switched regulatory text that appears in some official documents—a direction for "
        "future expansion. Fourth, the human IAA evaluation covers 60 of the 406 items; "
        "extending human agreement measurement to the full benchmark would provide stronger "
        "statistical guarantees, though the current sample is representative across all four "
        "task types and difficulty levels."
    )

    # ---- 8. Conclusion ----
    section_heading("8. Conclusion")
    body(
        "We have introduced IndiaFinBench, the first publicly available evaluation benchmark "
        "for LLM performance on Indian financial regulatory text. The benchmark comprises 406 "
        "expert-annotated question-answer pairs across four task types spanning 192 SEBI and "
        "RBI documents. Evaluating eleven contemporary models on the full benchmark reveals a "
        "clear tier structure: a top group clustering around 81–90% overall accuracy, a middle "
        "group around 75–79%, and one clear underperformer at 70%. Paired bootstrap "
        "significance testing establishes which differences are statistically robust."
    )
    body(
        "Key findings include: Gemini 2.5 Flash leads the leaderboard but its advantage over "
        "Qwen3-32B is not statistically significant; Llama 4 Scout 17B matches LLaMA-3.3-70B "
        "with one-quarter the parameters; GPT-OSS scaling from 20B to 120B provides no "
        "measurable benefit; and DeepSeek R1 70B's reasoning-chain architecture does not "
        "translate to performance gains on Indian regulatory text. Across all models, numerical "
        "reasoning and temporal reasoning emerge as the hardest tasks."
    )
    body(
        "IndiaFinBench highlights the importance of geographically and jurisdictionally diverse "
        "evaluation benchmarks. Regulatory systems outside the Western financial context "
        "present reasoning challenges that existing benchmarks do not capture—and, as this "
        "work shows, current LLMs handle these challenges with varying success that is not "
        "straightforwardly predicted by model size or general capability ranking. We make the "
        "full dataset, evaluation harness, all model predictions, and figure generation code "
        "publicly available to support ongoing research in multilingual and domain-specific "
        "financial NLP."
    )

    # ---- Ethics ----
    section_heading("Ethics Statement")
    body(
        "IndiaFinBench is constructed entirely from publicly available primary source documents "
        "released by the Securities and Exchange Board of India (sebi.gov.in) and the Reserve "
        "Bank of India (rbi.org.in). These documents are published by the Government of India "
        "for public use and carry no copyright restrictions on research use. No personally "
        "identifiable information is present in any source document or derived annotation. "
        "The benchmark is designed to evaluate model performance on regulatory reasoning tasks "
        "and does not contain any toxic, harmful, or privacy-violating content. The dataset is "
        "released under CC BY 4.0 to enable open research use with attribution."
    )

    # ---- Acknowledgements ----
    section_heading("Acknowledgements")
    body(
        "The author thanks the annotators who contributed to secondary validation and the "
        "human inter-annotator agreement evaluation. Evaluation infrastructure used the Groq "
        "API, Google AI Studio, Anthropic API, Moonshot AI (Kimi K2), and OpenRouter (GPT-OSS "
        "models). Local model inference used Ollama. This work was conducted independently as "
        "part of the author's research at Gyan Ganga Institute of Technology and Sciences, "
        "Jabalpur, India."
    )

    # ---- References ----
    section_heading("References")
    refs = [
        "Chen, Z., et al. (2021). FinQA: A Dataset of Numerical Reasoning over Financial "
        "Data. EMNLP 2021.",
        "Chalkidis, I., et al. (2022). LexGLUE: A Benchmark Dataset for Legal Language "
        "Understanding in English. ACL 2022.",
        "Dua, D., et al. (2019). DROP: A Reading Comprehension Benchmark Requiring Discrete "
        "Reasoning Over Paragraphs. NAACL 2019.",
        "Hendrycks, D., et al. (2021). CUAD: An Expert-Annotated NLP Dataset for Legal "
        "Contract Review. NeurIPS 2021.",
        "Hendrycks, D., et al. (2021). Measuring Massive Multitask Language Understanding. "
        "ICLR 2021.",
        "Islam, S., et al. (2023). FinanceBench: A New Benchmark for Financial Question "
        "Answering. arXiv:2311.11944.",
        "Landis, J.R., & Koch, G.G. (1977). The measurement of observer agreement for "
        "categorical data. Biometrics, 33(1), 159–174.",
        "Liang, P., et al. (2022). Holistic Evaluation of Language Models. NeurIPS 2022 "
        "(HELM).",
        "Loukas, L., et al. (2022). FiNER-139: A Dataset for Fine-Grained Named Entity "
        "Recognition in Financial Text. ACL 2022.",
        "Malik, V., et al. (2021). ILDC for CJPE: Indian Legal Documents Corpus for Court "
        "Judgment Prediction and Explanation. ACL 2021.",
        "Rajpurkar, P., et al. (2016). SQuAD: 100,000+ Questions for Machine Comprehension "
        "of Text. EMNLP 2016.",
        "Shah, A., et al. (2022). FLUE: Financial Language Understanding Evaluation. "
        "EMNLP 2022.",
        "Wilson, E.B. (1927). Probable inference, the law of succession, and statistical "
        "inference. Journal of the American Statistical Association, 22(158), 209–212.",
        "Zheng, Z., et al. (2022). ConvFinQA: Exploring the Chain of Numerical Reasoning "
        "in Conversational Finance Question Answering. EMNLP 2022.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"; r.font.size = Pt(11)

    # ---- Appendix ----
    doc.add_page_break()
    section_heading("Appendix A: Source Document Categories")
    body(
        "SEBI Documents: SEBI (Issue of Capital and Disclosure Requirements) Regulations 2018, "
        "SEBI (Listing Obligations and Disclosure Requirements) Regulations 2015, SEBI "
        "(Substantial Acquisition of Shares and Takeovers) Regulations 2011, SEBI (Prohibition "
        "of Insider Trading) Regulations 2015, SEBI (Alternative Investment Funds) Regulations "
        "2012, SEBI (Portfolio Managers) Regulations 2020, SEBI (Research Analysts) "
        "Regulations 2014, SEBI (Buy-Back of Securities) Regulations 2018, SEBI (Delisting of "
        "Equity Shares) Regulations 2021, SEBI (Mutual Funds) Regulations 1996, SEBI "
        "(Depositories and Participants) Regulations 2018, SEBI (Merchant Bankers) Regulations "
        "1992, recent SEBI circulars (2024–2026)."
    )
    body(
        "RBI Documents: RBI Monetary Policy Statements (2024–2026), RBI Master Directions on "
        "Unique Identifiers in Financial Markets, RBI (Small Finance Banks—Prudential Norms on "
        "Declaration of Dividend) Directions 2026, Government Securities auction notifications "
        "(2025–2026), State Government securities auction press releases, RBI Weekly "
        "Statistical Supplement extracts, RBI circulars on KYC/AML compliance."
    )

    section_heading("Appendix B: System Prompt Template")
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.5)
    p_code.paragraph_format.space_after = Pt(6)
    r_code = p_code.add_run(
        "You are an expert in Indian financial regulation and policy. Answer questions\n"
        "using ONLY the provided context passage. Do not use any external knowledge.\n"
        "Be concise and precise. Give only the answer — no preamble."
    )
    r_code.font.name = "Courier New"; r_code.font.size = Pt(10)

    # Save
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build_doc()
