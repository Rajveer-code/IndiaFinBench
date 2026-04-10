"""
update_paper_v7_fixed.py
Applies all 12 changes to produce indiafinbench_paper_v7_final.docx from v6.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import csv, json
import numpy as np
from pathlib import Path
from collections import defaultdict
from statsmodels.stats.proportion import proportion_confint
import docx
from docx.oxml.ns import qn
from copy import deepcopy

# ── Paths ────────────────────────────────────────────────────────────────────
PAPER_V6   = Path("paper/indiafinbench_paper_v6_final.docx")
PAPER_V7   = Path("paper/indiafinbench_paper_v7_final.docx")
RESULTS    = Path("evaluation/results")

# ── Model manifest (CSV file → label) ────────────────────────────────────────
MODELS = [
    ("haiku_results.csv",          "Claude 3 Haiku"),
    ("gemini_results.csv",         "Gemini 2.5 Flash"),
    ("llama4scout_results.csv",    "Llama 4 Scout 17B"),
    ("qwen3_32b_results.csv",      "Qwen3-32B"),
    ("groq70b_results.csv",        "LLaMA-3.3-70B"),
    ("deepseek_r1_70b_results.csv","DeepSeek R1 70B"),
    ("gemma4_e4b_results.csv",     "Gemma 4 E4B"),
    ("llama3_results.csv",         "LLaMA-3-8B"),
    ("mistral_results.csv",        "Mistral-7B"),
]

TASK_MAP = {
    "regulatory_interpretation": "REG",
    "numerical_reasoning":       "NUM",
    "contradiction_detection":   "CON",
    "temporal_reasoning":        "TMP",
}

# ── Load & compute stats ──────────────────────────────────────────────────────
def load_csv(path):
    with open(path, encoding="utf-8") as f:
        return list(csv.DictReader(f))

def get_stats(rows):
    s = {t: [] for t in ["REG","NUM","CON","TMP","overall"]}
    d = {"easy": [], "medium": [], "hard": []}
    for r in rows:
        t = TASK_MAP.get(r["task_type"])
        if not t:
            continue
        c = int(r["correct"])
        s[t].append(c); s["overall"].append(c)
        diff = r.get("difficulty", "").lower()
        if diff in d:
            d[diff].append(c)
    s["diff"] = d
    return s

def acc(lst):
    return sum(lst) / len(lst) * 100 if lst else 0.0

model_rows = {}
stats      = {}
for fname, label in MODELS:
    p = RESULTS / fname
    if not p.exists():
        print(f"MISSING: {p}"); continue
    rows = load_csv(p)
    model_rows[label] = rows
    stats[label]      = get_stats(rows)

sorted_labels = sorted(stats, key=lambda m: acc(stats[m]["overall"]), reverse=True)

# ── Print verification table ──────────────────────────────────────────────────
print(f"\n{'Model':<25} {'REG':>7} {'NUM':>7} {'CON':>7} {'TMP':>7} {'Overall':>9}")
print("─" * 62)
for lbl in sorted_labels:
    s = stats[lbl]
    print(f"{lbl:<25} {acc(s['REG']):>7.1f} {acc(s['NUM']):>7.1f} "
          f"{acc(s['CON']):>7.1f} {acc(s['TMP']):>7.1f} {acc(s['overall']):>9.1f}")
print("─" * 62)
print(f"{'Human Expert (n=30)':<25} {'55.6':>7} {'44.4':>7} {'83.3':>7} {'66.7':>7} {'60.0':>9}")

overall_accs = [acc(stats[m]["overall"]) for m in stats]
min_overall  = min(overall_accs)
max_overall  = max(overall_accs)
min_model    = min(stats, key=lambda m: acc(stats[m]["overall"]))
max_model    = max(stats, key=lambda m: acc(stats[m]["overall"]))
print(f"\nRange: {min_overall:.1f}% ({min_model}) → {max_overall:.1f}% ({max_model})\n")

# ── Helper: safe paragraph text replacement across runs ───────────────────────
def replace_text_in_para(para, old, new):
    """Replace `old` with `new` preserving run formatting where possible."""
    if old not in para.text:
        return False
    # Try within a single run first
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Text spans runs — fall back to clear + single run (loses inline bold)
    full = para.text.replace(old, new)
    style_name = para.style.name
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = full
    else:
        para.add_run(full)
    return True

def set_cell(cell, text):
    """Set a table cell's text, clearing existing content."""
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(text)

# ── Load document ─────────────────────────────────────────────────────────────
doc = docx.Document(str(PAPER_V6))

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 1 — Abstract: "five" → "nine", update range
# ═══════════════════════════════════════════════════════════════════════════════
ch1_done = False
for p in doc.paragraphs:
    if "five contemporary LLMs" in p.text:
        replace_text_in_para(p, "five contemporary LLMs", "nine contemporary LLMs")
        # Update range if present
        replace_text_in_para(p, "72.7% (Mistral-7B) to 91.3% (Claude 3 Haiku)",
                             f"{min_overall:.1f}% ({min_model}) to {max_overall:.1f}% ({max_model})")
        replace_text_in_para(p, "72.7% to 91.3%", f"{min_overall:.1f}% to {max_overall:.1f}%")
        replace_text_in_para(p, "five model tiers", "nine models across four tiers")
        ch1_done = True
# Also handle List Number para that contains the phrase
for p in doc.paragraphs:
    if "five contemporary LLMs" in p.text:
        replace_text_in_para(p, "five contemporary LLMs", "nine contemporary LLMs")
        replace_text_in_para(p, "three model tiers and two deployment modalities",
                             "four model tiers and two deployment modalities")
        ch1_done = True

print(f"CHANGE 1 (Abstract): {'✅' if ch1_done else '❌ not found'}")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 7 / 4.1 Text: Update model grouping narrative
# ═══════════════════════════════════════════════════════════════════════════════
ch7_done = False
for p in doc.paragraphs:
    if "We evaluate five models spanning a range" in p.text:
        p.clear()
        p.add_run(
            "We evaluate nine contemporary models spanning a range of sizes and providers, "
            "grouped as: (1) Frontier API models — Claude 3 Haiku (Anthropic) and Gemini 2.5 Flash (Google); "
            "(2) Reasoning model — DeepSeek R1 70B (DeepSeek, via OpenRouter); "
            "(3) Large open-weight API models — Llama 4 Scout 17B (Meta, via Groq), "
            "Qwen3-32B (Alibaba, via Groq), and LLaMA-3.3-70B (Meta, via Groq); "
            "(4) Small local models — Gemma 4 E4B (Google, via Ollama), "
            "LLaMA-3-8B (Meta, via Ollama), and Mistral-7B (Mistral AI, via Ollama)."
        )
        ch7_done = True
        break
print(f"CHANGE 7 (4.1 model grouping narrative): {'✅' if ch7_done else '❌ not found'}")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 12 — Deprecation notes in Section 4.1 para [63]
# ═══════════════════════════════════════════════════════════════════════════════
ch12_done = False
for p in doc.paragraphs:
    if "claude-3-haiku-20240307" in p.text and "All API-based evaluations" in p.text:
        note = (
            " Note: claude-3-haiku-20240307 was retired by Anthropic on April 19, 2026, "
            "after evaluation was completed. Results are fully reproducible using cached "
            "per-item outputs released with the dataset. "
            "deepseek-r1-distill-llama-70b was retired from Groq on October 2, 2025 and "
            "was evaluated via OpenRouter, which hosts identical model weights."
        )
        if p.runs:
            p.runs[-1].text += note
        else:
            p.add_run(note)
        ch12_done = True
        break
if not ch12_done:
    # Try appending to any para with that identifier
    for p in doc.paragraphs:
        if "claude-3-haiku-20240307" in p.text:
            p.add_run(
                " Note: claude-3-haiku-20240307 was retired by Anthropic on April 19, 2026, "
                "after evaluation was completed. Results are fully reproducible using cached "
                "per-item outputs released with the dataset."
            )
            ch12_done = True
            break
print(f"CHANGE 12 (Deprecation notes): {'✅' if ch12_done else '❌ not found'}")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 5.1 narrative update ("all five models" → "all nine models")
# ═══════════════════════════════════════════════════════════════════════════════
for p in doc.paragraphs:
    replace_text_in_para(p, "all five models", "all nine models")
    replace_text_in_para(p, "five models spanning", "nine models spanning")
    replace_text_in_para(p, "average model accuracy of 81.7%", "average model accuracy of 79.4%")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 6 — Section 5.1 Narrative: Add DeepSeek/Gemma analytical paragraph
# ═══════════════════════════════════════════════════════════════════════════════
gemma_overall = acc(stats["Gemma 4 E4B"]["overall"])
gemma_reg     = acc(stats["Gemma 4 E4B"]["REG"])
gemma_tmp     = acc(stats["Gemma 4 E4B"]["TMP"])

ch6_done = False
for i, p in enumerate(doc.paragraphs):
    if "Table 1: IndiaFinBench Results" in p.text:
        # Insert new analytical paragraph BEFORE the table caption
        new_p = p.insert_paragraph_before(
            "DeepSeek R1 70B achieves 70.7% overall — the lowest among all models — "
            "yet exhibits the second-highest CON score (93.3%), surpassed only by Llama 4 Scout 17B (100.0%). "
            "Its REG accuracy (60.4%) and TMP accuracy (60.0%) both fall below several non-reasoning models, "
            "suggesting that explicit chain-of-thought reasoning aids binary contradiction detection but does "
            "not compensate for domain knowledge gaps on Indian regulatory interpretation and temporal reasoning. "
            f"Gemma 4 E4B ({gemma_overall:.1f}% overall) performs competitively on REG ({gemma_reg:.1f}%) "
            f"but struggles on TMP ({gemma_tmp:.1f}%), placing it between LLaMA-3-8B (75.3%) and "
            "Mistral-7B (72.7%) in the local-model tier. "
            "These findings indicate that reasoning capability alone is insufficient for Indian regulatory text — "
            "domain grounding and temporal amendment-chain knowledge matter independently of general reasoning capacity."
        )
        ch6_done = True
        break
print(f"CHANGE 6 (5.1 new analytical paragraph): {'✅' if ch6_done else '❌ not found'}")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 10 — Section 7.1 Discussion: Add reasoning-model paragraph
# ═══════════════════════════════════════════════════════════════════════════════
ch10_done = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("IndiaFinBench is intentionally challenging by design"):
        # Insert reasoning-model discussion paragraph BEFORE this one
        new_p = p.insert_paragraph_before(
            "Reasoning vs. Non-Reasoning Models. DeepSeek R1 70B's strong CON performance (93.3%) "
            "aligns with the hypothesis that explicit chain-of-thought reasoning aids binary judgment tasks, "
            "where the answer requires a clear yes/no decision over an explicitly provided pair of statements. "
            "However, its underperformance on REG (60.4%) and TMP (60.0%) — both below several non-reasoning "
            "models including LLaMA-3-8B (77.4% REG, 74.3% TMP) — suggests that chain-of-thought reasoning "
            "without domain-specific training data does not reliably improve performance on Indian regulatory text. "
            "Regulatory interpretation requires precise recall of thresholds, definitions, and procedural rules "
            "that are not derivable by reasoning alone; temporal reasoning in this domain requires tracking "
            "amendment chains across regulatory notifications, a retrieval-heavy task not addressed by longer "
            "reasoning traces. This implies that future work targeting this domain should prioritise "
            "domain-adaptive pretraining or retrieval-augmented generation over reasoning-chain prompting alone."
        )
        ch10_done = True
        break
print(f"CHANGE 10 (7.1 reasoning-model discussion): {'✅' if ch10_done else '❌ not found'}")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 11 — Conclusion: Replace "five contemporary models", update range
# ═══════════════════════════════════════════════════════════════════════════════
ch11_done = False
for p in doc.paragraphs:
    if "five contemporary models" in p.text and "Conclusion" not in p.text:
        replace_text_in_para(p, "five contemporary models", "nine contemporary models")
        replace_text_in_para(p, "performance ranges from 72.7% to 91.3%",
                             f"performance ranges from {min_overall:.1f}% to {max_overall:.1f}%")
        replace_text_in_para(p, "72.7% to 91.3%", f"{min_overall:.1f}% to {max_overall:.1f}%")
        # Add DeepSeek finding if not already mentioned
        if "DeepSeek" not in p.text:
            p.add_run(
                " A notable new finding is that DeepSeek R1 70B, despite being a dedicated reasoning model, "
                f"achieves only {min_overall:.1f}% overall, with strong CON (93.3%) but weak REG (60.4%) "
                "and TMP (60.0%), demonstrating that reasoning capability alone does not substitute for "
                "domain grounding in Indian regulatory text."
            )
        ch11_done = True
print(f"CHANGE 11 (Conclusion): {'✅' if ch11_done else '❌ not found'}")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 2 — Models Table [3]: Add Llama 4 Scout, Qwen3, DeepSeek R1, Gemma 4
# ═══════════════════════════════════════════════════════════════════════════════
tbl_models = doc.tables[3]
existing_model_names = {row.cells[0].text.strip() for row in tbl_models.rows[1:]}

new_model_rows = [
    ("Llama 4 Scout 17B",  "Meta (via Groq)",          "17B",         "API",   "meta-llama/llama-4-scout-17b-16e-instruct (Groq), accessed March 30, 2026"),
    ("Qwen3-32B",          "Alibaba (via Groq)",         "32B",         "API",   "qwen3-32b (Groq), accessed March 30, 2026"),
    ("DeepSeek R1 70B",    "DeepSeek (via OpenRouter)",  "70B distilled","API",  "deepseek/deepseek-r1-distill-llama-70b (OpenRouter), accessed April 10, 2026. Note: retired from Groq on October 2, 2025; evaluated via OpenRouter, which hosts identical model weights."),
    ("Gemma 4 E4B",        "Google (via Ollama)",        "4B",           "Local","gemma4 (Ollama local), accessed April 10, 2026"),
]

added_to_models = 0
for (name, provider, params, access, identifier) in new_model_rows:
    if name not in existing_model_names:
        row = tbl_models.add_row()
        set_cell(row.cells[0], name)
        set_cell(row.cells[1], provider)
        set_cell(row.cells[2], params)
        set_cell(row.cells[3], access)
        set_cell(row.cells[4], identifier)
        added_to_models += 1

print(f"CHANGE 2 (Models table): ✅ added {added_to_models} new rows "
      f"(total now {len(tbl_models.rows)-1} models)")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 3 — Results Table [4]: Rebuild with all 9 models
# ═══════════════════════════════════════════════════════════════════════════════
tbl_res = doc.tables[4]

# Preserve special rows (Human Expert, Random Baseline)
special_rows = []
for row in tbl_res.rows[1:]:
    first_cell = row.cells[0].text.strip()
    if first_cell.startswith("Human Expert") or first_cell.startswith("Random Baseline"):
        special_rows.append([c.text for c in row.cells])

# Remove all non-header rows
while len(tbl_res.rows) > 1:
    tbl_res._element.remove(tbl_res.rows[-1]._element)

# Add sorted model rows
for lbl in sorted_labels:
    s = stats[lbl]
    row = tbl_res.add_row()
    set_cell(row.cells[0], lbl)
    set_cell(row.cells[1], f"{acc(s['REG']):.1f}")
    set_cell(row.cells[2], f"{acc(s['NUM']):.1f}")
    set_cell(row.cells[3], f"{acc(s['CON']):.1f}")
    set_cell(row.cells[4], f"{acc(s['TMP']):.1f}")
    set_cell(row.cells[5], f"{acc(s['overall']):.1f}")

# Average row
avg_row = tbl_res.add_row()
set_cell(avg_row.cells[0], "Average")
for i, task in enumerate(["REG","NUM","CON","TMP","overall"]):
    avg_val = np.mean([acc(stats[m][task]) for m in sorted_labels])
    set_cell(avg_row.cells[i+1], f"{avg_val:.1f}")

# Re-add special rows
for cells in special_rows:
    row = tbl_res.add_row()
    for j, val in enumerate(cells):
        set_cell(row.cells[j], val)

print(f"CHANGE 3 (Results table): ✅ {len(sorted_labels)} model rows + Average + "
      f"{len(special_rows)} special rows")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 4 — Difficulty Table [5]: Rebuild with all 9 models
# ═══════════════════════════════════════════════════════════════════════════════
tbl_diff = doc.tables[5]
while len(tbl_diff.rows) > 1:
    tbl_diff._element.remove(tbl_diff.rows[-1]._element)

for lbl in sorted_labels:
    d = stats[lbl]["diff"]
    row = tbl_diff.add_row()
    set_cell(row.cells[0], lbl)
    set_cell(row.cells[1], f"{acc(d['easy']):.1f}")
    set_cell(row.cells[2], f"{acc(d['medium']):.1f}")
    set_cell(row.cells[3], f"{acc(d['hard']):.1f}" if d['hard'] else "—")

print(f"CHANGE 4 (Difficulty table): ✅ {len(sorted_labels)} model rows")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 5 — Error Table [6]: Rebuild with all 9 models
# ═══════════════════════════════════════════════════════════════════════════════
def compute_errors(rows, overall_acc):
    nrf = trf = dkf = cgf = 0
    for r in rows:
        if r["correct"] == "1":
            continue
        t = r["task_type"]
        if t == "numerical_reasoning":
            nrf += 1
        elif t == "temporal_reasoning":
            trf += 1
        elif t == "regulatory_interpretation":
            dkf += 1
        elif t == "contradiction_detection":
            cgf += 1
    tot = nrf + trf + dkf + cgf
    return nrf, trf, dkf, cgf, tot

tbl_err = doc.tables[6]
while len(tbl_err.rows) > 1:
    tbl_err._element.remove(tbl_err.rows[-1]._element)

for lbl in sorted_labels:
    overall_score = acc(stats[lbl]["overall"])
    nrf, trf, dkf, cgf, tot = compute_errors(model_rows[lbl], overall_score)
    if tot == 0:
        continue
    row = tbl_err.add_row()
    set_cell(row.cells[0], lbl)
    set_cell(row.cells[1], f"{dkf} ({dkf/tot*100:.0f}%)")
    set_cell(row.cells[2], f"{nrf} ({nrf/tot*100:.0f}%)")
    set_cell(row.cells[3], f"{trf} ({trf/tot*100:.0f}%)")
    set_cell(row.cells[4], f"{cgf} ({cgf/tot*100:.0f}%)")
    # Total Errors column
    if len(row.cells) > 5:
        set_cell(row.cells[5], str(tot))

print(f"CHANGE 5 (Error table): ✅ {len(sorted_labels)} model rows")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 8 — Wilson CI Table [8]: Rebuild with all 9 models
# ═══════════════════════════════════════════════════════════════════════════════
def wilson_ci(correct_list, n_list):
    k = sum(correct_list)
    n = len(n_list)
    lo, hi = proportion_confint(k, n, alpha=0.05, method="wilson")
    return f"[{lo*100:.1f}%, {hi*100:.1f}%]"

tbl_ci = doc.tables[8]
while len(tbl_ci.rows) > 1:
    tbl_ci._element.remove(tbl_ci.rows[-1]._element)

for lbl in sorted_labels:
    s = stats[lbl]
    row = tbl_ci.add_row()
    set_cell(row.cells[0], lbl)
    for i, task in enumerate(["REG","NUM","CON","TMP"]):
        lst = s[task]
        k = sum(lst); n = len(lst)
        lo, hi = proportion_confint(k, n, alpha=0.05, method="wilson")
        set_cell(row.cells[i+1], f"[{lo*100:.1f}%, {hi*100:.1f}%]")

print(f"CHANGE 8 (Wilson CI table): ✅ {len(sorted_labels)} model rows")

# ═══════════════════════════════════════════════════════════════════════════════
# CHANGE 9 — Bootstrap Table [9]: Add new pairs for DeepSeek R1 & Gemma 4
# ═══════════════════════════════════════════════════════════════════════════════
def paired_bootstrap(a, b, n_resamples=10_000, seed=42):
    np.random.seed(seed)
    diffs = np.array(a, dtype=float) - np.array(b, dtype=float)
    n = len(diffs)
    boot = np.random.choice(diffs, size=(n_resamples, n), replace=True).mean(axis=1)
    p    = np.mean(np.abs(boot) >= np.abs(diffs.mean()))
    return float(p), float(diffs.mean() * 100)

tbl_boot = doc.tables[9]

# Existing pairs (already in table)
existing_pairs = set()
for row in tbl_boot.rows[1:]:
    existing_pairs.add((row.cells[0].text.strip(), row.cells[1].text.strip()))

new_models    = ["DeepSeek R1 70B", "Gemma 4 E4B"]
other_models  = [m for m in sorted_labels if m not in new_models]
pairs_added   = 0

for new_m in new_models:
    if new_m not in stats:
        continue
    a = stats[new_m]["overall"]
    for other_m in other_models:
        if (new_m, other_m) in existing_pairs or (other_m, new_m) in existing_pairs:
            continue
        b = stats[other_m]["overall"]
        pval, delta = paired_bootstrap(a, b)
        row = tbl_boot.add_row()
        set_cell(row.cells[0], new_m)
        set_cell(row.cells[1], other_m)
        sign = "+" if delta > 0 else ""
        set_cell(row.cells[2], f"{sign}{delta:.1f}pp")
        p_str = f"{pval:.4f}" if pval >= 0.0001 else "<0.0001"
        if pval < 0.05:
            p_str += " *"
        set_cell(row.cells[3], p_str)
        pairs_added += 1

print(f"CHANGE 9 (Bootstrap table): ✅ {pairs_added} new pairs added")

# ═══════════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(str(PAPER_V7))
print(f"\n✅ Saved: {PAPER_V7}")
print("\nSummary:")
print("  CHANGE  1 — Abstract 'five'→'nine', range updated")
print("  CHANGE  2 — Models table: 4 new rows added (Llama4Scout, Qwen3, DeepSeek, Gemma4)")
print("  CHANGE  3 — Results table rebuilt with all 9 models, sorted by accuracy")
print("  CHANGE  4 — Difficulty table rebuilt with all 9 models")
print("  CHANGE  5 — Error table rebuilt with all 9 models")
print("  CHANGE  6 — 5.1 new DeepSeek/Gemma analytical paragraph inserted")
print("  CHANGE  7 — 4.1 model grouping narrative updated (4 groups)")
print("  CHANGE  8 — Wilson CI table rebuilt with all 9 models")
print("  CHANGE  9 — Bootstrap table: new pairs for DeepSeek R1 & Gemma 4")
print("  CHANGE 10 — 7.1 reasoning-model discussion paragraph added")
print("  CHANGE 11 — Conclusion updated: 'nine', new range, DeepSeek insight")
print("  CHANGE 12 — Deprecation notes for Claude 3 Haiku & DeepSeek R1")
